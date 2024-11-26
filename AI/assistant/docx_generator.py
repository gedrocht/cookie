import re
from docx import Document
from docx.shared import Inches
import sys

sys.path.append('../')
from utils import util

def log(msg):
    util.log(msg, "DOCXGN")

def generate_document(filename, text):
    # Create a new Word document
    doc = Document()

    # Regex patterns for Markdown formatting
    bold_pattern = r'\*\*(.*?)\*\*'
    italic_pattern = r'_(.*?)_'

    # Start a single paragraph
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.33)

    """ Adds text with bold and italic formatting to the provided paragraph """
    # Split the text by bold formatting and apply styling
    parts = re.split(bold_pattern, text)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Add non-bold text, split further by italic formatting
            italic_parts = re.split(italic_pattern, part)
            for j, italic_part in enumerate(italic_parts):
                if j % 2 == 0:
                    # Normal text
                    paragraph.add_run(italic_part)
                else:
                    # Italic text
                    run = paragraph.add_run(italic_part)
                    run.italic = True
        else:
            # Bold text
            run = paragraph.add_run(part)
            run.bold = True
    # Save the document
    filename = filename.replace(",","").replace(".","")
    doc.save(filename + ".docx")
    log(f"Generated {filename}.docx")